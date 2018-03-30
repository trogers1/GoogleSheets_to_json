#from __future__ import print_function
import httplib2
import os
from sys import exit

from apiclient import discovery
from oauth2client import client
from oauth2client import tools
from oauth2client.file import Storage

#for creating a word doc:
from docx import Document

try:
    import argparse
    flags = argparse.ArgumentParser(parents=[tools.argparser]).parse_args()
except ImportError:
    flags = None

# If modifying these scopes, delete your previously saved credentials
# at ~/.credentials/sheets.googleapis.com-python-quickstart.json
SCOPES = 'https://www.googleapis.com/auth/spreadsheets.readonly'
CLIENT_SECRET_FILE = 'client_secret_quickstart.json'
APPLICATION_NAME = 'Google Sheets API Python Quickstart'

document = Document()

def get_credentials():
    """Gets valid user credentials from storage.

    If nothing has been stored, or if the stored credentials are invalid,
    the OAuth2 flow is completed to obtain the new credentials.

    Returns:
        Credentials, the obtained credential.
    """
    home_dir = os.path.expanduser('~')
    credential_dir = os.path.join(home_dir, '.credentials')
    if not os.path.exists(credential_dir):
        os.makedirs(credential_dir)
    credential_path = os.path.join(credential_dir,
                                   'sheets.googleapis.com-python-quickstart.json')

    store = Storage(credential_path)
    credentials = store.get()
    if not credentials or credentials.invalid:
        flow = client.flow_from_clientsecrets(CLIENT_SECRET_FILE, SCOPES)
        flow.user_agent = APPLICATION_NAME
        if flags:
            credentials = tools.run_flow(flow, store, flags)
        else: # Needed only for compatibility with Python 2.6
            credentials = tools.run(flow, store)
        print 'Storing credentials to ' + credential_path
    return credentials


def main():
    """Shows basic usage of the Sheets API.

    Creates a Sheets API service object and prints the names and majors of
    students in a sample spreadsheet:
    https://docs.google.com/spreadsheets/d/1_gC-Oj31LiUtH6fjCi7uh7_26hGeLIevphvCg7I-_EU/edit
    """
    credentials = get_credentials()
    http = credentials.authorize(httplib2.Http())
    discoveryUrl = ('https://sheets.googleapis.com/$discovery/rest?'
                    'version=v4')
    service = discovery.build('sheets', 'v4', http=http,
                              discoveryServiceUrl=discoveryUrl)
    get_sheet_id(service)


def get_sheet_id(service):
    print '''
        Let's find your speadsheet. Can you give me your spreadsheet ID?
        It should look something like this: '1_gC-Oj31LiUtH6fjCi7uh7_26hGeLIevphvCg7I-_EU'
        It can be copied from the shareable url of the spreadsheet.

        NOTE: Remember you can quit the program at any time by pressing Ctrl-C.
        '''
    spreadsheetId = raw_input('> ')
    if len(str(spreadsheetId)) > 46 or len(str(spreadsheetId)) < 43:
        print '''
            This doesn't look right. Most spreadsheet IDs are 44 characters in length.
            Yours is '''+str(len(str(spreadsheetId)))+''' characters long.
            Continue with this ID or try again?'''
        answer = raw_input('> ')
        if 'C' in answer or 'c' in answer or 'continue' in answer or "Continue" in answer:
            get_sheet_name(service, spreadsheetId)
        if 'try' in answer or 'again' in answer:
            get_sheet_id(service)
    else:
        get_sheet_name(service, spreadsheetId)
#    spreadsheetId = '1_gC-Oj31LiUtH6fjCi7uh7_26hGeLIevphvCg7I-_EU'


def get_sheet_name(service, spreadsheetId):
    print '''
        Alright, now what is the name of your sheet?
        The default name for a sheet is 'Sheet1'.
        If there is a space in the name, wrap the name in single quotes (i.e. 'Sheet One').
        '''
    sheetName = raw_input('> ')
#    sheetName = 'Sheet1'
    get_cell(service, spreadsheetId, sheetName)

def get_cell(service, spreadsheetId, sheetName):
    print '''
        Ok, so we are working in '''+str(sheetName)+'''.

        What is the cell you want me to place in the new document?
        You should type a capital letter and a number--
        Ex: A1, E5, F13, etc.
        '''
    cellID = raw_input('> ')

    rangeName = sheetName+'!'+cellID
    print rangeName
    #rangeName = 'Sheet1!A1' #this is the sheet/tab name and cell range in A1 notation
    result = service.spreadsheets().values().get(
        spreadsheetId=spreadsheetId, range=rangeName).execute()
    values = result.get('values', [])

    if not values:
        print '''
        No data found.
        Please choose a different cell.
            '''
        get_cell(service, spreadsheetId, sheetName)
    else:
        print '''
        This is what I found:
            '''
        print '''
            %s
            ''' % values
        print '''
        Is the data inside the brackets and single quotes correct?
            (Y / N)
            '''
        correct = raw_input('> ')
        if 'y' in correct or 'Y' in correct:
            isHeading(service, values, spreadsheetId, sheetName)
        elif 'n' in correct or 'N' in correct:
            get_cell(service, spreadsheetId, sheetName)
        else:
            print '''
        Sorry I didn't understand that.
                '''
            get_cell(service, spreadsheetId, sheetName)


def isHeading(service, values, spreadsheetId, sheetName):
        print '''
        Is this a Heading? (Y / N)
        '''
        choice = raw_input('> ')
        if 'Y' in choice or "y" in choice or "yes" in choice or "Yes" in choice:
            create_Heading(service, values, spreadsheetId, sheetName)
        elif 'N' in choice or "n" in choice or "no" in choice or "No" in choice:
            create_Paragraph(service, values, spreadsheetId, sheetName)
        else:
            print '''
        Sorry I didn't understand that.
                    '''
            isHeading(service, values, spreadsheetId, sheetName)

def create_Heading(service, values, spreadsheetId, sheetName):
        print '''
        Ok! Heading it is.
        '''
        heading = document.add_heading(str(values), level=1)
        for paragraph in document.paragraphs or document.heading:
            if '[[u\'' in paragraph.text or heading.text:
                print 'Cleaning text...'
                htext = heading.text.replace('[[u\'', '')
                heading.text = htext
        for paragraph in document.paragraphs:
            if '\']]' in paragraph.text:
                print 'Cleaning text...'
                htext = heading.text.replace('\']]', '')
                heading.text = htext
        done(service, values, spreadsheetId, sheetName)


def create_Paragraph(service, values, spreadsheetId, sheetName):
        print '''
        Ok! It\'s not a heading.
        '''
        paragraph = document.add_paragraph(str(values))
        for paragraph in document.paragraphs:
            if '[[u\'' in paragraph.text:
                print 'Cleaning text...'
                ptext = paragraph.text.replace('[[u\'', '')
                paragraph.text = ptext
        for paragraph in document.paragraphs:
            if '\']]' in paragraph.text:
                print 'Cleaning text...'
                ptext = paragraph.text.replace('\']]', '')
                paragraph.text = ptext
        done(service, values, spreadsheetId, sheetName)

def done(service, values, spreadsheetId, sheetName):
        print '''
        Now that that is done, would you like to add anything else to the
        document?
        (Y / N)
            '''
        choice = raw_input('> ')
        if 'Y' in choice or "y" in choice or "yes" in choice or "Yes" in choice:
            main2(service, spreadsheetId, sheetName)
        elif 'N' in choice or "n" in choice or "no" in choice or "No" in choice:
            saveDoc()
        else:
            print '''
        Sorry I didn't understand that.
                    '''
            done(service, values, spreadsheetId, sheetName)


def main2(service, spreadsheetId, sheetName):
    print'''
        Are we still using the same Google Sheets document?
        (Y / N)
        '''
    choice = raw_input('> ')
    if 'Y' in choice or "y" in choice or "yes" in choice or "Yes" in choice:
        sameSheet(service, spreadsheetId, sheetName)
    elif 'N' in choice or "n" in choice or "no" in choice or "No" in choice:
        diffDocument(service, sheetName)
    else:
        print '''
        Sorry I didn't understand that.
                '''
        main2(service, spreadsheetId, sheetName)

def sameSheet(service, spreadsheetId, sheetName):
    print '''
        Are we still using the same sheet name of the document?
        We are currently using '''+str(sheetName)+'''.
        (Y / N)
        '''
    choice = raw_input('> ')
    if 'Y' in choice or "y" in choice or "yes" in choice or "Yes" in choice:
        get_cell(service, spreadsheetId, sheetName)
    elif 'N' in choice or "n" in choice or "no" in choice or "No" in choice:
        get_sheet_name(service, spreadsheetId)
    else:
        print '''
        Sorry I didn't understand that.
                '''
        sameSheet(service, spreadsheetId, sheetName)

def diffDocument(service, sheetName):
    print '''
        What\'s the new spreadsheet ID?
        Remember, it should look something like this:
        '1_gC-Oj31LiUtH6fjCi7uh7_26hGeLIevphvCg7I-_EU'
        It can be copied from the shareable url of the spreadsheet.
        '''
    spreadsheetId = raw_input('> ')
    if len(str(spreadsheetId)) > 46 or len(str(spreadsheetId)) < 43:
        print '''
            This doesn't look right. Most spreadsheet IDs are 44 characters in length.
            Yours is '''+str(len(str(spreadsheetId)))+''' characters long.
            Continue with this ID or try again?'''
        answer = raw_input('> ')
        if 'C' in answer or 'c' in answer or 'continue' in answer or "Continue" in answer:
            sameSheet(service, spreadsheetId, sheetName)
        if 'try' in answer or 'again' in answer:
            diffDocument(service, sheetName)
    else:
        sameSheet(service, spreadsheetId, sheetName)




def saveDoc():
    print '''
        Alright. Let's save the document. What would you like it to be called?
        '''
    saveName = raw_input('> ')
    print '''
        Great. We\'ll call it '''+str(saveName)+'''. Where would you like it
        saved? Paste a path to a file/directory of your choice.
        If you want to keep the current working directory, type "wd".
        '''
    savePlace = raw_input('> ')
    if 'wd' in savePlace:
        document.save(str(saveName)+'.docx')
        print 'Document created in working directory.'
        exit(0)
    else:
        document.save(str(savePlace)+'/'+str(saveName)+'.docx')
        print 'Document created in '+str(savePlace)+'/'+str(saveName)
        exit(0)

main()
